import os # Manipulación de rutas y archivos
import tempfile # Creación de directorios temporales para almacenar archivos de forma aislada
import zipfile # Empaquetamiento de archivos en formato ZIP
import pandas as pd # Lectura y manipulación de datos (En este caso desde archivos excel)
from docx import Document # Se trabaja con la manipulación de documentos Word
from docx.shared import Pt # Definición de tamaños en puntos para fuentes
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Alineación de parrafos
from docx.oxml.ns import qn # Manejo de nombres de espacio en XML para compatibilidad en word
from datetime import datetime # Manejo de fechas y horas
import shutil # Operaciones de alto nivel con archivos (copiar, mover, etc...)
import sys # Acceso a variables y funciones del sistema, como los argumentos del script
import locale # Configuración de la localización (idioma, formato de fecha, etc...)

# Directorio para almacenar los documentos generados como historial (registro permanente)
HISTORIAL_DIR = os.path.join(os.getcwd(), "reports_historial")
if not os.path.exists(HISTORIAL_DIR):
    os.makedirs(HISTORIAL_DIR) # Se crea el directorio en caso de no existir

#=================================================================================================
# Función: apply_styles
# Objetivo: Aplicar estilos predefinidos a un fragmento de texto (run) dentro del documento word.
#=================================================================================================
def apply_styles(run, font_size=9):
    run.font.name = 'Arial' # Fuente del texto
    run.font.size = Pt(font_size) # Tamaño del texto en puntos
    run.font.bold = True # Texto en negritas activado
    r = run._element
    # Se establece la fuente para la compatibilidad con versiones de Word que usan codificación diferente para ciertos idiomas
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

#=======================================================================================================
# Función: replace_mark
# Objetivo: Reemplazar marcadores (placeholders) en el documento Word, tanto en párrafos como en tablas,
#           y aplicar los estilos definidos a esos textos.
#=======================================================================================================
def replace_mark(doc, mark, text):
    try:
        text = str(text) # Asegura que el contenido a insertar sea una cadenass
        # Se itera sobre cada párrafo del documento 
        for paragraph in doc.paragraphs:
            if mark in paragraph.text:
                # Se itera sobre cada fragmento de texto (run) dentro del párrafo
                for run in paragraph.runs:
                    if mark in run.text:
                        # Se reemplaza el marcador con el texto designado
                        run.text = run.text.replace(mark, str(text))
                        # Se aplican los estilos personalizados al run modificado
                        apply_styles(run)
    
        # Se itera sobre cada tabla del documento
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if mark in cell.text: # Se verifica si el marcador existe en la celda
                        for paragraph in cell.paragraphs: # Se recorre cada párrafo dentro de la celda
                            for run in paragraph.runs: # Recorre cada fragmento de texto (run) del párrafo
                                if mark in run.text:
                                    run.text = run.text.replace(mark, str(text))
                                    # Se aplican los estilos personalizados
                                    apply_styles(run)

                        # Se justifica el texto de cada celda al centro
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(f"Error al reemplazar el marcador {mark}: {e}")

#=======================================================================================================
# Función: add_participants
# Objetivo: Agregar los datos de los participantes a una tabla en el documento Word.
#           Dado el formato proporcionado se asume que la tabla de los participantes es la tercera del
#           documento.
#=======================================================================================================
def add_participants(doc, participants):
    try:
        table = doc.tables[2] # Se asume que la tabla de los participantes es la tercera del documento
    except IndexError:
        raise IndexError("Error: No se encontró la tabla esperada de los participantes, revisar el formato Word")

    num_participants = len(participants) # Número total de participantes a agregar
    total_rows = max(10, num_participants) # Asegura al menos 10 filas, incluso si hay menos participantes

    # Se itera para crear filas en la tabla
    for idx in range(total_rows):
        # Se agrega una nueva fila a la tabla
        new_row = table.add_row()

        if idx < num_participants:
            participant = participants[idx]
            # Se preparan los datos: Índice, RPE y nombre completo del participante
            data = [str(idx + 1), participant['RPE'], participant['NOMBRE_COMPLETO']]
            # Se obtiene el dato del sexo, si está presente
            sex = participant.get('SEXO_TRAB', '')
            # Dependiendo del valor de 'SEXO_TRAB', se marca la celda correspondiente
            if sex == 'M':
                data.append('X') # Celda del sexo masculino
                data.append('') # Celda del sexo femenino vacía
            elif sex == 'F':
                data.append('') # Celda del sexo masculino vacía
                data.append('X') # Celda del sexo femenino
            else:
                data.append('')
                data.append('') # Datos vacíos si no hay información
        else:
            data = ['', '', '', '', ''] # Filas vacías de relleno si no existen más participantes

        # Se recorre cada dato y se coloca en la celda correspondiente
        for i, text in enumerate(data):
            # Se coloca cada dato correspondiente en su respectiva celda
            cell = new_row.cells[i]
            cell.text = text
            # Se accede al parrafo dentro de la celda actual
            paragraph = cell.paragraphs[0]
            # Se aplica la justificación al centro para el texto
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Se aplican los estilos con un tamaño de fuente reducido para ajustarse al formato de la tabla
            if paragraph.runs:
                apply_styles(paragraph.runs[0], font_size=7.5)

#=======================================================================================================
# Función: generate_reports
# Objetivo: Generar los reportes a partir de los datos de un archivo Excel,
#           reemplazar marcadores en una plantilla Word, agregar participantes,
#           y empaquetar los reportes generados en un archivo ZIP.
#=======================================================================================================
def generate_reports():
    try:
        # Ruta del archivo Excel con la información de cursos y participantes
        excel_path = 'db_excel.xlsx'
        if not os.path.exists(excel_path):
            error_message = f"Error: No se encontró el archivo {excel_path}. Verifique su existencia o ubicación"
            print(error_message, flush=True)
            sys.exit(1)

        # Lectura de la hoja 'P01' que contiene los cursos
        df_courses = pd.read_excel(excel_path, sheet_name='P01')
        # print(df_courses.columns) ---- Se imprimen las columnas de la hoja de los cursos
        # Lectura de la hoja 'PARTIP01' que contiene a los participantes
        df_participants = pd.read_excel(excel_path, sheet_name='PARTIP01')
        # print(df_participants.columns) ---- Se imprimen las columnas de la hoja de los participantes
    except ValueError as e:
        error_message = f"Error: El archivo {excel_path} no contiene alguna de las hojas esperadas. Detalles: {e}"
        print(error_message, flush=True)
        sys.exit(1)
    except PermissionError as e:
        error_message = f"Error: No se puede acceder al archivo {excel_path}. Puede que este abierto por otro programa o este bloqueado. Cierre el archivo e intente nuevamente: {e}"
        print(error_message, flush=True)
        sys.exit(1)
    except Exception as e:
        print(f"Error inesperado al abrir el documento excel: {e}")
        return

    # Determina el mes a procesar: se usa el argumento pasado al script o, si no, el mes actual
    if len(sys.argv) > 1 and sys.argv[1]:
        try:
            selected_month = int(sys.argv[1]) # Argumento pasado: Mes seleccionado por el usuario
        except ValueError:
            selected_month = datetime.now().month
    else:
        selected_month = datetime.now().month
        
    # Se obtiene el mes seleccionado o el actual (1 = Enero, 2 = Febrero, etc...)
    current_month = selected_month
    # Se configura el locale para obtener el nombre del mes en español
    locale.setlocale(locale.LC_TIME, 'es_ES.utf8')
    # Se obtiene el nombre del mes en español una vez ya configurado el locale
    # Se crea un objeto datetime con un día cualquiera para ese mes y se formatea el nombre (capitalizado)
    month_name = datetime(1900, current_month, 1).strftime('%B').capitalize()

    # Verifica que las columnas esenciales existan en ambos Dataframes
    required_columns = {'ID_CURSO', 'MES_PROGRAMADO', 'FECHA_INICIO', 'FECHA_TERMINO', 'ID_ACTIVIDAD'}
    if not required_columns.issubset(df_courses.columns) or not required_columns.issubset(df_participants.columns):
        print("Error: El archivo o algunas de las hojas no contiene las columnas necesarias: ID_CURSO, MES_PROGRAMADO, FECHA_INICIO, FECHA_TERMINO")
        exit()

    # Se unen los datos de cursos y participantes basándose en 'ID_CURSO'
    df_complete = pd.merge(df_courses, df_participants, on='ID_CURSO', how='left')
    # Se renombran las columnas duplicadas para su posterior uso a la hora de filtrar a los participantes por mes y fechas
    df_complete.rename(columns={'MES_PROGRAMADO_x': 'MES_PROGRAMADO_COUR', 'MES_PROGRAMADO_y': 'MES_PROGRAMADO_PART'}, inplace=True)
    df_complete.rename(columns={'FECHA_INICIO_x': 'FECHA_INICIO_COUR', 'FECHA_INICIO_y': 'FECHA_INICIO_PART'}, inplace=True)
    df_complete.rename(columns={'FECHA_TERMINO_x': 'FECHA_TERMINO_COUR', 'FECHA_TERMINO_y': 'FECHA_TERMINO_PART'}, inplace=True)
    df_complete.rename(columns={'ID_ACTIVIDAD_x': 'ID_ACTIVIDAD_COUR', 'ID_ACTIVIDAD_y': 'ID_ACTIVIDAD_PART'}, inplace=True)
    # print(df_complete.columns) ---- Se imprimen las columnas del Dataframe resultante de la unión de la tabla de los cursos y los participantes

    # Filtra los cursos que corresponden al mes a procesar
    df_filtered = df_courses.loc[df_courses['MES_PROGRAMADO'] == current_month]

    # Inicialización de contadores para reportar el resultado final
    total_docs_generated = 0
    courses_without_participants = 0

    # Se crea una lista para almacenar las rutas de los reportes generados
    reports_generated = []
    zip_filename = None # Se inicializa la variable para evitar UnboundLocalError (Cuando se usa una variable antes de asignarle un valor)

    if df_filtered.empty:
        print("No hay cursos disponibles para este mes")
    else:
        # Se crea un directorio temporal donde se guardarán los reportes generados
        with tempfile.TemporaryDirectory() as temp_dir:
            print(f"Directorio temporal creado: {temp_dir}")

            # Se itera sobre cada curso filtrado
            for index, row in df_filtered.iterrows():
                # Se obtienen los participantes del curso actual y así tambien solo los que esten en el curso del mes actual junto con sus fechas correspondientes
                participants = df_complete.loc[(df_complete['ID_CURSO'] == row['ID_CURSO']) 
                                               & (df_complete['MES_PROGRAMADO_PART'] == current_month)
                                               & (df_complete['FECHA_INICIO_PART'] == row['FECHA_INICIO']) 
                                               & (df_complete['FECHA_TERMINO_PART'] == row['FECHA_TERMINO'])
                                               & (df_complete['ID_ACTIVIDAD_PART'] == row['ID_ACTIVIDAD'])]
                
                # Convierte la información de los participantes a una lista de diccionarios y elimina duplicados
                participants_list = participants[['RPE', 'NOMBRE_COMPLETO', 'SEXO_TRAB']].fillna('').drop_duplicates(subset=['RPE', 'NOMBRE_COMPLETO']).to_dict('records')
                # print(participants_list) ---- Se imprime la lista de diccionarios de los participantes

                # Se calcula el número de lotes los cuales se dividen por 10 participantes en cada uno
                num_batches = (len(participants_list) // 10) + (1 if len(participants_list) % 10 else 0)
                # Se itera sobre cada lote
                for batch in range(num_batches):
                    try:
                        # Se carga la plantilla de Word para el reporte
                        doc = Document('FORMATO DE LISTA DE ASISTENCIA enero 2025.docx')
                    except FileNotFoundError:
                        print("Error: No se encontró el documento base de Word.")
                        exit()
                    except Exception as e:
                        print(f"Error al abrir el documento Word: {e}")
                        exit()

                    try:
                        # Se reemplazan los marcadores en la plantilla por los datos del curso
                        replace_mark(doc, "[NOMBRE_CURSO]", row['NOMBRE_CURSO'])
                        replace_mark(doc, "[FECHA_INICIO]", row['FECHA_INICIO'])
                        replace_mark(doc, "[FECHA_TERMINO]", row['FECHA_TERMINO'])
                    except Exception as e:
                        print(f"Error al reemplazar los marcadores en el documentos Word: {e}")
                        continue

                    # Se agregan los participantes a la tabla del documento
                    if participants_list:
                        batch_participants = participants_list[batch * 10:(batch + 1) * 10]
                        try:
                            add_participants(doc, batch_participants)
                        except Exception as e:
                            print(f"Error al agregar a los participantes en el documento Word: {e}")
                            continue
                    else:
                        print(f"No hay participantes inscritos en el curso {row['NOMBRE_CURSO']}")
                        courses_without_participants += 1

                    # Genera un nombre de archivo seguro para el reporte y lo guarda en el directorio temporal
                    file_name = os.path.join(temp_dir, f'{row["NOMBRE_CURSO"].replace("/", "_").replace(" ", "_")}_'f'{row["FECHA_INICIO"]}_'f'{row["FECHA_TERMINO"]}_L{batch + 1}_'f'{row["ID_ACTIVIDAD"]}.docx')
                    try:    
                        doc.save(file_name)
                        # print(f"Documento generado: {file_name}") ---- Se imprime el documento generado en la iteración
                        reports_generated.append(file_name)
                        total_docs_generated += 1
                        # Se copia el documento al historial (queda registrado permanentemente)
                        historial_path = os.path.join(HISTORIAL_DIR, os.path.basename(file_name))
                        shutil.copy(file_name, historial_path)
                    except Exception as e:
                        print(f"Error al guardar el documento: {e}")

            try:
                if not reports_generated:
                    print(reports_generated)
                    error_message = "No existen documentos generados para comprimir"
                    print(error_message, flush=True)
                    sys.exit(1)
                # Se empaquetan todos los reportes generados en un archivo ZIP
                zip_filename = os.path.abspath(os.path.join(os.getcwd(), f"Reportes_{month_name}.zip"))
                with zipfile.ZipFile(zip_filename, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for report in reports_generated:
                        try:
                            zipf.write(report, arcname=os.path.basename(report))
                            # print(f"Agregado al ZIP: {os.path.basename(report)}") ---- Imprime los archivos que se almacenan en el .zip
                        except FileNotFoundError:
                            print(f"Error: No se encontró el archivo {report}. No sera incluido en el ZIP.")
                        except Exception as e:
                            print(f"Error al agregar {report} al zip: {e}")
                print("Archivo ZIP creado exitosamente");
            except ValueError as ve:
                print(f"Advertencia: {ve}")
            except Exception as e:
                print(f"Error inesperado al crear el archivo ZIP: {e}")

            # Mensaje final de resumen
            print("\n--- Resumen de generación de reportes ---")
            print(f"Mes procesado: {month_name}")
            print(f"Total de documentos generados: {total_docs_generated}")
            print(f"Cursos sin participantes: {courses_without_participants}")

            return zip_filename

if __name__ == '__main__':
    # print("Working directory:", os.getcwd()) ---- Imprime el directorio actual en el que se trabaja

    zip_file = generate_reports()
    if zip_file:
        print(f"ZIP generado: {zip_file}") # Ruta completa de la ubicación del .zip
